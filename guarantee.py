import docx
from docx.shared import *
# from docx.shared import RGBColor
# from docx.shared import Cm, Pt  #加入可調整的 word 單位
# from docx.shared import RGBColor
# from docx.shared import Cm, Pt  #加入可調整的 word 單位
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH #處理字串的置中
import win32com.client as win32
import os
# from docx import Document

# point specific file
folder_path = "C:/Users/KevinChen/Desktop/保證金"

def guarantee_content(date, customer_name, customer_id, amount):
    # Create a word file
    doc = docx.Document()

    # Document content
    doc.add_heading('EVOX(易喂)服務保證金繳交證明', level=1)
    doc.add_paragraph('\n日期:  中華民國  '+ date[0:3] +'  年  '+ date[3:5] +' 月 '+ date[5:] +' 日 ')
    doc.add_paragraph('客戶名稱 : '+ customer_name + ' (統編: '+ customer_id +') ')
    doc.add_paragraph('茲收到客戶繳交之EVOX(易喂)服務保證金共計新台幣____' + amount  +'____元，關於此保證金之歸還，請詳見客戶與本公司簽署之EVOX(易喂)服務申請書附件之EVOX服務契約條款第二條第7、8項。 ')
    doc.add_paragraph('易得雲端股份有限公司 \n統編: 90616421 \n地址: 台北市114內湖區內湖路一段316號5樓之1 ')

    # adjust the form of content 
    for paragraph in doc.paragraphs:
        paragraph.style.font.size = Pt(12)

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            paragraph.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER # LEFT, CENTER, RIGHT
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.name = '標楷體'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
                run.font.size = Pt(20)
            
    doc.add_picture("圖片1.png", width = Cm(10))

    # store
    word_file_path = os.path.join(folder_path, customer_name +'.docx') # 指定word儲存路徑

    doc.save(word_file_path) # 儲存word檔

    return word_file_path

def transfer_word_to_pdf(word_file_path):

    # transfer word file to pdf
    pdf_file_path = os.path.join(folder_path, customer_name +'.pdf') #指定PDF儲存路徑
    wdFormatPDF = 17
    word_app = win32.gencache.EnsureDispatch('Word.Application')
    doc = word_app.Documents.Open(word_file_path)
    doc.SaveAs(pdf_file_path, FileFormat=wdFormatPDF)
    doc.Close()
    word_app.Quit()

# main 
while True:
    date = input('日期 1120101:')
    customer_name = input('客戶名稱: ')
    customer_id = input('客戶統編: ')
    if len(customer_id) != 8:
        while True:
            customer_id = input('客戶統編: ')
            if len(customer_id) == 8:
                break
    amount = input('保證金金額: ')

    word_file_path = guarantee_content(date, customer_name, customer_id, amount)

    transfer_word_to_pdf(word_file_path)
